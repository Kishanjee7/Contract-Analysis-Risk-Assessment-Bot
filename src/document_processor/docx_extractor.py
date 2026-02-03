"""
DOCX Text Extractor using python-docx
"""
from docx import Document
from pathlib import Path
from typing import Optional, Dict, List
import io


class DOCXExtractor:
    """Extract text content from DOCX files"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def extract(self, file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> Dict:
        """
        Extract text from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            file_bytes: Raw bytes of the DOCX file
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            if file_bytes:
                doc_source = io.BytesIO(file_bytes)
            elif file_path:
                doc_source = file_path
            else:
                raise ValueError("Either file_path or file_bytes must be provided")
            
            doc = Document(doc_source)
            
            paragraphs = []
            sections = []
            current_section = {"heading": "", "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"heading": text, "content": []}
                else:
                    paragraphs.append(text)
                    current_section["content"].append(text)
            
            # Add last section
            if current_section["content"] or current_section["heading"]:
                sections.append(current_section)
            
            # Extract metadata
            metadata = {
                "num_paragraphs": len(paragraphs),
                "num_sections": len(sections)
            }
            
            if doc.core_properties:
                metadata.update({
                    "author": doc.core_properties.author or "",
                    "title": doc.core_properties.title or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                })
            
            full_text = "\n\n".join(paragraphs)
            
            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "sections": sections,
                "metadata": metadata,
                "format": "docx"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "format": "docx"
            }
    
    def extract_tables(self, file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> List[Dict]:
        """
        Extract tables from a DOCX file
        
        Returns:
            List of tables found in the document
        """
        try:
            if file_bytes:
                doc_source = io.BytesIO(file_bytes)
            elif file_path:
                doc_source = file_path
            else:
                return []
            
            doc = Document(doc_source)
            tables = []
            
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "table_num": i + 1,
                    "data": table_data
                })
            
            return tables
            
        except Exception as e:
            return []
